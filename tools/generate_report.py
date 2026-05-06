from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(r"C:\rating-work-git")
SCREENSHOTS_DIR = Path(r"C:\rating-work-screenshots")
REPORT_DIR = PROJECT_ROOT / "report"

STUDENT_FULL_NAME = "Грузин Михаил Дмитриевич"
GROUP = "о.ИЗДтс 23.1/Б3-24"
TEACHER = "Королькова Марина Николаевна"
CITY_YEAR = "Москва, 2026"
DISCIPLINE = "Проектная деятельность в ИТ"
PROJECT_NAME = "personal-agile-planner"
GITHUB_USERNAME = "mikhail-lab1"
VARIANT = "Разработка программы планировщика Agile для личного пользования"


def report_file_name(full_name: str) -> str:
    parts = full_name.split()
    if len(parts) >= 3:
        surname, name, patronymic = parts[:3]
        return f"{surname}_{name[0]}.{patronymic[0]}_git.docx"
    return f"{full_name.replace(' ', '_')}_git.docx"


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(14)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_centered_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    if level == 1:
        document.add_page_break()
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)


def screenshots() -> Iterable[Path]:
    if not SCREENSHOTS_DIR.exists():
        return []
    patterns = ("*.png", "*.jpg", "*.jpeg")
    result: list[Path] = []
    for pattern in patterns:
        result.extend(SCREENSHOTS_DIR.glob(pattern))
    return sorted(result, key=lambda item: item.name.lower())


def caption_from_name(path: Path) -> str:
    name = path.stem.replace("_", " ").replace("-", " ")
    return " ".join(part.capitalize() if not part.isdigit() else part for part in name.split())


def build_report() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    add_page_number(document.sections[0])

    add_centered_paragraph(document, "Кафедра информационных систем")
    document.add_paragraph()
    add_centered_paragraph(document, "Рейтинговая работа", bold=True)
    add_centered_paragraph(document, "Расчетно-аналитическое задание")
    add_centered_paragraph(document, f"по дисциплине: {DISCIPLINE}")
    add_centered_paragraph(document, "Задание / Вариант: индивидуальное задание")
    add_centered_paragraph(document, f"Тема: {VARIANT}")
    document.add_paragraph()

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    title_data = [
        ("Выполнена обучающимся группы", GROUP),
        ("ФИО обучающегося", STUDENT_FULL_NAME),
        ("Преподаватель", TEACHER),
        ("Проект", PROJECT_NAME),
    ]
    for row_index, (left, right) in enumerate(title_data):
        set_cell_text(table.rows[row_index].cells[0], left)
        set_cell_text(table.rows[row_index].cells[1], right)

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    add_centered_paragraph(document, CITY_YEAR)

    add_heading(document, "ОГЛАВЛЕНИЕ")
    toc_items = [
        "Введение",
        "1. Основания для разработки и назначение разработки",
        "2. Создание и настройка Git/GitHub",
        "3. ИСР и жизненный цикл информационной системы",
        "4. Настройка локального и удаленного репозитория",
        "Заключение",
        "Список использованных источников",
        "Приложения",
    ]
    for item in toc_items:
        add_body_paragraph(document, item)

    add_heading(document, "ВВЕДЕНИЕ")
    add_body_paragraph(
        document,
        "Рейтинговая работа посвящена созданию учебного ИТ-проекта и демонстрации "
        "применения системы контроля версий Git при разработке программного продукта. "
        f"Вариант выбран по первой букве фамилии обучающегося: {VARIANT}.",
    )
    add_body_paragraph(
        document,
        "Цель работы - разработать простую консольную программу Personal Agile Planner, "
        "создать структуру проекта, настроить локальный Git-репозиторий, оформить ветки "
        "по иерархической структуре работ и подготовить материалы для публикации в GitHub.",
    )

    add_heading(document, "1. ОСНОВАНИЯ ДЛЯ РАЗРАБОТКИ И НАЗНАЧЕНИЕ РАЗРАБОТКИ")
    add_body_paragraph(
        document,
        "Основанием для разработки является задание рейтинговой работы по дисциплине "
        f'"{DISCIPLINE}". Разрабатываемая программа предназначена для личного планирования '
        "задач с использованием простых элементов Agile-подхода.",
    )
    add_body_paragraph(
        document,
        "Пользователь программы может вести список задач, задавать приоритет, срок выполнения "
        "и отслеживать статус работы. Данные сохраняются локально в JSON-файле, что делает "
        "проект простым для изучения и проверки.",
    )

    add_heading(document, "2. СОЗДАНИЕ И НАСТРОЙКА GIT/GITHUB")
    add_body_paragraph(
        document,
        "Для проекта была создана рабочая папка C:\\rating-work-git. В ней подготовлены "
        "исходные файлы программы, документация, тесты и служебные сценарии автоматизации. "
        "Локальный репозиторий создается командой git init, после чего файлы добавляются "
        "в индекс командой git add и фиксируются командой git commit.",
    )
    add_bullet(document, "git config используется для проверки имени и email автора коммитов.")
    add_bullet(document, "git status показывает состояние файлов рабочей области и индекса.")
    add_bullet(document, "git branch и git switch используются для работы с ветками.")
    add_bullet(document, "git stash демонстрирует временное сохранение незавершенных изменений.")
    add_bullet(document, "git log используется для просмотра истории коммитов.")

    add_heading(document, "3. ИСР И ЖИЗНЕННЫЙ ЦИКЛ ИНФОРМАЦИОННОЙ СИСТЕМЫ")
    add_body_paragraph(
        document,
        "Для проекта используется каскадная модель жизненного цикла с последовательными "
        "этапами: инициация, требования, проектирование, реализация, тестирование и релиз. "
        "Каждому этапу соответствует отдельная ветка Git.",
    )
    add_table(
        document,
        ["Этап", "Ветка Git", "Содержание работ"],
        [
            ["Стабильная версия", "main", "Проверенная версия проекта"],
            ["Разработка", "develop", "Объединение текущих изменений"],
            ["Инициация", "wbs/01-initiation", "Определение темы, цели и назначения"],
            ["Требования", "wbs/02-requirements", "Описание требований"],
            ["Проектирование", "wbs/03-design", "Структура программы и формат данных"],
            ["Реализация", "wbs/04-implementation", "Разработка модулей программы"],
            ["Тестирование", "wbs/05-testing", "Проверка запуска и тестов"],
            ["Релиз", "wbs/06-release", "Подготовка публикации и отчета"],
        ],
    )

    add_heading(document, "4. НАСТРОЙКА ЛОКАЛЬНОГО И УДАЛЕННОГО РЕПОЗИТОРИЯ")
    add_body_paragraph(
        document,
        "Локальный репозиторий хранит историю разработки проекта. После отдельного подтверждения "
        "пользователя к нему может быть подключен публичный GitHub-репозиторий с помощью команды "
        "git remote add origin. Отправка изменений выполняется командами git push. Для проверки "
        "удаленного репозитория используются git clone, git fetch и git pull.",
    )
    add_body_paragraph(
        document,
        f"Планируемая учетная запись GitHub: {GITHUB_USERNAME}. Публикация в GitHub не выполняется "
        "автоматически и требует отдельного подтверждения.",
    )

    add_heading(document, "ЗАКЛЮЧЕНИЕ")
    add_body_paragraph(
        document,
        "В ходе работы подготовлен учебный проект Personal Agile Planner, создана структура "
        "исходного кода, добавлена проектная документация и настроен локальный Git-репозиторий. "
        "Ветки проекта отражают этапы иерархической структуры работ и жизненного цикла "
        "информационной системы.",
    )

    add_heading(document, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "Git Documentation [Электронный ресурс]. URL: https://git-scm.com/doc",
        "GitHub Docs [Электронный ресурс]. URL: https://docs.github.com/",
        "Python Documentation [Электронный ресурс]. URL: https://docs.python.org/3/",
        "Методические указания по выполнению рейтинговой работы по дисциплине "
        '"Проектная деятельность в ИТ".',
    ]
    for index, source in enumerate(sources, start=1):
        add_body_paragraph(document, f"{index}. {source}")

    add_heading(document, "ПРИЛОЖЕНИЯ")
    image_paths = list(screenshots())
    if not image_paths:
        add_body_paragraph(
            document,
            "Скриншоты не обнаружены. Поместите реальные PNG/JPG-файлы в папку "
            f"{SCREENSHOTS_DIR} и запустите генератор отчета повторно.",
        )
    else:
        for index, image_path in enumerate(image_paths, start=1):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Cm(15.5))

            caption = document.add_paragraph(
                f"Рисунок {index} - {caption_from_name(image_path)}"
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.first_line_indent = Cm(0)

    output_path = REPORT_DIR / report_file_name(STUDENT_FULL_NAME)
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_report()
    print(f"Report created: {path}")

